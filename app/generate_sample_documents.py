import os
import logging
from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import docx
import re

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

def create_pdf_from_markdown(markdown_path, output_path):
    """
    Create a PDF file from a markdown text file.
    
    Args:
        markdown_path: Path to the markdown file
        output_path: Path to save the PDF file
    """
    try:
        # Read markdown content
        with open(markdown_path, 'r') as file:
            markdown_text = file.read()
        
        # Convert Path object to string if needed
        if not isinstance(output_path, str):
            output_path = str(output_path)
        
        # Create PDF
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        flowables = []
        
        # Process markdown content
        lines = markdown_text.split('\n')
        
        for line in lines:
            style = styles['Normal']
            
            # Headers
            if line.startswith('# '):
                style = styles['Title']
                line = line[2:]
            elif line.startswith('## '):
                style = styles['Heading1']
                line = line[3:]
            elif line.startswith('### '):
                style = styles['Heading2']
                line = line[4:]
                
            # Skip code blocks and empty lines in this simple converter
            if line.startswith('```') or not line.strip():
                continue
                
            para = Paragraph(line, style)
            flowables.append(para)
            flowables.append(Spacer(1, 6))
            
        doc.build(flowables)
        logging.info(f"Created PDF: {output_path}")
        return True
        
    except Exception as e:
        logging.error(f"Error creating PDF: {e}")
        return False

def create_docx_from_markdown(markdown_path, output_path):
    """
    Create a DOCX file from a markdown text file.
    
    Args:
        markdown_path: Path to the markdown file
        output_path: Path to save the DOCX file
    """
    try:
        # Read markdown content
        with open(markdown_path, 'r') as file:
            markdown_text = file.read()
        
        # Convert Path object to string if needed
        if not isinstance(output_path, str):
            output_path = str(output_path)
            
        # Create DOCX
        doc = docx.Document()
        
        # Process markdown content
        lines = markdown_text.split('\n')
        in_code_block = False
        
        for line in lines:
            # Handle code blocks
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue
                
            if in_code_block:
                doc.add_paragraph(line, style='No Spacing')
                continue
            
            # Headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=0)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
            # Lists (simplified)
            elif line.strip().startswith('- '):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            # Normal paragraph
            elif line.strip():
                doc.add_paragraph(line)
                
        doc.save(output_path)
        logging.info(f"Created DOCX: {output_path}")
        return True
        
    except Exception as e:
        logging.error(f"Error creating DOCX: {e}")
        return False

def main():
    # Ensure data directory exists
    data_dir = Path(__file__).parent.parent / 'data'
    data_dir.mkdir(exist_ok=True)
    
    # PDF sample
    pdf_content_path = data_dir / 'sample_pdf_content.txt'
    pdf_output_path = data_dir / 'ecommerce_technical_docs.pdf'
    
    if not pdf_content_path.exists():
        logging.error(f"PDF content file not found: {pdf_content_path}")
    else:
        create_pdf_from_markdown(pdf_content_path, pdf_output_path)
    
    # DOCX sample
    doc_content_path = data_dir / 'sample_doc_content.txt'
    doc_output_path = data_dir / 'ecommerce_returns_policy.docx'
    
    if not doc_content_path.exists():
        logging.error(f"DOCX content file not found: {doc_content_path}")
    else:
        create_docx_from_markdown(doc_content_path, doc_output_path)
    
    print("\nSample documents created successfully!")
    print(f"PDF: {pdf_output_path}")
    print(f"DOCX: {doc_output_path}")
    print("\nTo process these documents and create embeddings, run:")
    print(f"python app/insert_document_vectors.py setup")
    print(f"python app/insert_document_vectors.py pdf {pdf_output_path}")
    print("\nTo search the embeddings, run:")
    print(f"python app/document_search.py interactive")

if __name__ == "__main__":
    main()